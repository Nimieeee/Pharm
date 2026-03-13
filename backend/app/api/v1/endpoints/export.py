from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from uuid import UUID
from datetime import datetime
import io
import markdown
import re
import json

from app.core.database import get_db
from app.services.auth import AuthService
from app.services.chat import ChatService
from app.services.manuscript_formatter import ManuscriptFormatter
from app.api.v1.endpoints.ai import get_chat_service
from app.models.user import User
from supabase import Client

try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from xhtml2pdf import pisa
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

router = APIRouter()


async def get_export_user(
    request: Request,
    token: str = Query(None),
    db: Client = Depends(get_db)
) -> User:
    """Authenticate via Bearer header OR query param token (for GET downloads)."""
    auth_service = AuthService(db)
    auth_header = request.headers.get("Authorization", "")
    jwt_token = None
    if auth_header.startswith("Bearer "):
        jwt_token = auth_header[7:]
    elif token:
        jwt_token = token
    if not jwt_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        token_data = auth_service.verify_token(jwt_token)
        user = await auth_service.get_user_by_email(token_data.email)
        if not user or not user.is_active:
            raise HTTPException(status_code=401, detail="Not authenticated")
        return user
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Not authenticated")


import base64
import zlib
import httpx
from PIL import Image, ImageDraw, ImageFont

def encode_kroki(text: str) -> str:
    """Encode text for Kroki diagram rendering"""
    compressed = zlib.compress(text.encode('utf-8'), 9)
    return base64.urlsafe_b64encode(compressed).decode('utf-8')

def add_watermark(img_bytes: bytes, text="Benchside") -> bytes:
    """Add a watermark to the bottom right of the image"""
    try:
        img = Image.open(io.BytesIO(img_bytes)).convert("RGBA")
        
        # Make a blank image for the text, initialized to transparent text color
        txt = Image.new('RGBA', img.size, (255, 255, 255, 0))
        d = ImageDraw.Draw(txt)
        
        # Calculate position (bottom right)
        width, height = img.size
        margin = 10
        # Default font
        font = ImageFont.load_default()
        
        # Text bounding box
        bbox = d.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        x = width - text_width - margin - 5
        y = height - text_height - margin - 5
        
        # Draw text half-transparent blue
        d.text((x, y), text, fill=(0, 102, 204, 180), font=font)
        
        # Combine
        watermarked = Image.alpha_composite(img, txt)
        
        # Convert back to bytes
        out = io.BytesIO()
        watermarked.save(out, format="PNG")
        return out.getvalue()
    except Exception as e:
        print(f"Watermark error: {e}")
        return img_bytes


async def process_text_for_diagrams(text: str) -> list:
    """
    Extract mermaid diagrams, render them, and return a list of parts:
    either string text or a dict {"type": "image", "data": bytes}
    """
    parts = []
    
    # Simple regex to find ```mermaid \n ... \n``` blocks
    pattern = re.compile(r'```mermaid\n(.*?)\n```', re.DOTALL)
    
    last_end = 0
    
    # We'll use a sync approach inside an async function for simplicity, 
    # but in a real prod env we should use httpx.AsyncClient here.
    async with httpx.AsyncClient(timeout=10.0) as client:
        for match in pattern.finditer(text):
            # Add text before diagram
            if match.start() > last_end:
                parts.append({"type": "text", "content": text[last_end:match.start()]})
                
            mermaid_code = match.group(1).strip()
            
            # Inject a theme config to make colors readable
            theme_config = "%%\n%%{init: {'theme': 'base', 'themeVariables': { 'primaryColor': '#ffffff', 'primaryTextColor': '#333333', 'primaryBorderColor': '#0066cc', 'lineColor': '#666666', 'tertiaryColor': '#f4f4f4', 'tertiaryBorderColor': '#cccccc'}}}%%\n"
            full_code = theme_config + mermaid_code
            
            try:
                b64 = encode_kroki(full_code)
                resp = await client.get(f"https://kroki.io/mermaid/png/{b64}")
                if resp.status_code == 200:
                    img_bytes = add_watermark(resp.content)
                    parts.append({"type": "image", "data": img_bytes, "alt": "Diagram"})
                else:
                    parts.append({"type": "text", "content": f"\n[Diagram rendering failed: {resp.status_code}]\n"})
            except Exception as e:
                parts.append({"type": "text", "content": f"\n[Diagram generation error: {str(e)}]\n"})
                
            last_end = match.end()
            
    # Add remaining text
    if last_end < len(text):
        parts.append({"type": "text", "content": text[last_end:]})
        
    # If no diagrams found, just return the text
    if not parts:
        parts = [{"type": "text", "content": text}]
        
    return parts

def clean_markdown_for_docx(text: str) -> str:
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'#{1,6}\s', '', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text


@router.get("/conversations/{conversation_id}/export/docx")
async def export_conversation_docx(
    conversation_id: UUID,
    current_user: User = Depends(get_export_user),
    chat_service: ChatService = Depends(get_chat_service)
):
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx is not installed")

    conversation = await chat_service.get_conversation(conversation_id, current_user)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")

    messages = await chat_service.get_messages(conversation_id, current_user)

    doc = Document()
    title = doc.add_heading(conversation.title or "Chat Export", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("_" * 50)

    for msg in messages:
        p = doc.add_paragraph()
        if msg.role == 'user':
            runner = p.add_run("User:\n")
            runner.bold = True
            runner.font.color.rgb = RGBColor(0, 102, 204)
        else:
            runner = p.add_run("AI Assistant:\n")
            runner.bold = True
            runner.font.color.rgb = RGBColor(0, 153, 76)
            
        parts = await process_text_for_diagrams(msg.content)
        for part in parts:
            if part["type"] == "text":
                cleaned_text = clean_markdown_for_docx(part["content"])
                p.add_run(cleaned_text + "\n")
            elif part["type"] == "image":
                try:
                    p.add_run("\n")
                    # Add picture using a BytesIO stream
                    img_stream = io.BytesIO(part["data"])
                    doc.add_picture(img_stream, width=Inches(6.0)) # Scale to fit doc width
                    p = doc.add_paragraph() # Start a new paragraph after image
                except Exception as e:
                    p.add_run(f"\n[Failed to insert image: {str(e)}]\n")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"export_{conversation_id}.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/conversations/{conversation_id}/export/pdf")
async def export_conversation_pdf(
    conversation_id: UUID,
    current_user: User = Depends(get_export_user),
    chat_service: ChatService = Depends(get_chat_service)
):
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=500, detail="xhtml2pdf is not installed")

    conversation = await chat_service.get_conversation(conversation_id, current_user)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")

    messages = await chat_service.get_messages(conversation_id, current_user)

    html_content = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Helvetica, Arial, sans-serif; line-height: 1.6; color: #333; }}
            h1 {{ color: #2c3e50; text-align: center; border-bottom: 2px solid #eee; padding-bottom: 10px; }}
            .message {{ margin-bottom: 20px; padding: 15px; border-radius: 8px; }}
            .user {{ background-color: #f8f9fa; border-left: 4px solid #0066cc; }}
            .assistant {{ background-color: #ffffff; border-left: 4px solid #00994c; border: 1px solid #eee; }}
            .role {{ font-weight: bold; margin-bottom: 8px; font-size: 0.9em; text-transform: uppercase; letter-spacing: 0.5px; }}
            .user .role {{ color: #0066cc; }}
            .assistant .role {{ color: #00994c; }}
            .content {{ white-space: pre-wrap; }}
        </style>
    </head>
    <body>
        <h1>{conversation.title or 'Chat Export'}</h1>
        <div style="text-align: right; color: #666; margin-bottom: 30px;">
            Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        </div>
    """

    for msg in messages:
        role_class = 'user' if msg.role == 'user' else 'assistant'
        role_label = 'User' if msg.role == 'user' else 'AI Assistant'
        
        parts = await process_text_for_diagrams(msg.content)
        html_msg_content = ""
        
        for part in parts:
            if part["type"] == "text":
                safe_text = part["content"].replace('<', '&lt;').replace('>', '&gt;')
                if msg.role == 'assistant':
                    html_msg_content += markdown.markdown(safe_text, extensions=['fenced_code', 'tables'])
                else:
                    html_msg_content += f"<div class='content'>{safe_text}</div>"
            elif part["type"] == "image":
                # Embed image directly as base64 in HTML
                b64_img = base64.b64encode(part["data"]).decode('utf-8')
                html_msg_content += f'<div style="text-align: center; margin: 20px 0;"><img src="data:image/png;base64,{b64_img}" style="max-width: 100%; border: 1px solid #ddd; padding: 5px; border-radius: 4px;" alt="diagram"></div>'

        html_content += f"""
        <div class="message {role_class}">
            <div class="role">{role_label}</div>
            {html_msg_content}
        </div>
        """

    html_content += "</body></html>"

    file_stream = io.BytesIO()
    pisa_status = pisa.CreatePDF(html_content, dest=file_stream)

    if pisa_status.err:
        raise HTTPException(status_code=500, detail="PDF generation failed")

    file_stream.seek(0)

    filename = f"export_{conversation_id}.pdf"
    return StreamingResponse(
        file_stream,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/conversations/{conversation_id}/export/manuscript")
async def export_manuscript_docx(
    conversation_id: UUID,
    style: str = Query("report", regex="^(report|manuscript|plain)$"),
    current_user: User = Depends(get_export_user),
    chat_service: ChatService = Depends(get_chat_service)
):
    """
    Export conversation as a structured manuscript with sections.

    Styles:
    - report: Executive Summary, Introduction, Findings, Discussion, Conclusion
    - manuscript: Abstract, Introduction, Methods, Results, Discussion, References
    - plain: Simple formatted document with headings

    Features:
    - Title page with Benchside branding
    - Table of Contents (auto-generated in Word)
    - Numbered heading hierarchy
    - Page numbers in footer
    - 'Generated by Benchside' watermark
    """
    # Get messages
    messages = await chat_service.get_messages(conversation_id, current_user)

    if not messages:
        raise HTTPException(status_code=404, detail="No messages found")

    # Convert messages to dict format for formatter
    messages_dict = [
        {
            "role": msg.role,
            "content": msg.content,
            "created_at": msg.created_at.isoformat() if msg.created_at else None
        }
        for msg in messages
    ]

    # Create formatter with AI capability (use already-injected chat_service, not raw get_db())
    formatter = ManuscriptFormatter(
        multi_provider=chat_service.ai.multi_provider
        if hasattr(chat_service, 'ai') and hasattr(chat_service.ai, 'multi_provider')
        else None
    )

    try:
        # Structure content with AI
        structured = await formatter.structure_content(messages_dict, style=style)

        # Build DOCX
        docx_bytes = formatter.build_docx(structured, style=style)

        # Return as download
        file_stream = io.BytesIO(docx_bytes)
        file_stream.seek(0)

        filename = f"manuscript_{conversation_id}_{style}.docx"
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate manuscript: {str(e)}"
        )
