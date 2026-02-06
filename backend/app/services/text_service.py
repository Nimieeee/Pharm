from docx import Document
import io
import asyncio
from app.services.vision_service import _optimize_and_encode
from mistralai import Mistral
from app.utils.rate_limiter import mistral_limiter
import logging
try:
    from PIL import Image
except ImportError:
    Image = None

import httpx

logger = logging.getLogger(__name__)

# Constants for NVIDIA API
NVIDIA_API_URL = "https://integrate.api.nvidia.com/v1/chat/completions"
# Ideally this should be in .env, but user provided specific key for this switch. 
# We'll rely on the existing pattern or use the one provided if not in env.
# For now, we will expect it to be passed or hardcode the specific one requested if not present?
# The user wants to use *this* specific integration.
# We will check os.environ first, if not matching, we might need to set it.
# Actually, the user's snippet had a hardcoded key. I'll prefer env var passed in, 
# but for the "api_key" argument in this function, it usually comes from MISTRAL_API_KEY.
# We should probably add NVIDIA_API_KEY to settings/env. 
# For this refactor, I will add a fallback or use the passed key (which smart_loader passes).
# Wait, SmartLoader passes MISTRAL_API_KEY. 
# I will use a constant here for the NVIDIA key provided by user to ensure it works as requested,
# or better, update .env later. Let's use a variable.

NVIDIA_KEY = "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" 

async def process_text_document(content: bytes, ext: str, user_prompt: str, api_key: str = None):
    """
    Process text documents (DOCX, MD, TXT).
    For DOCX, it supports Multimodal extraction (Images -> Text) using NVIDIA Mistral Large 3.
    """
    text = ""
    try:
        if ext == 'docx':
            doc = Document(io.BytesIO(content))
            
            # 1. Extract Text & Tables
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    tables_text.append(" | ".join(row_data))
            
            text_content = "\n\n".join(paragraphs)
            if tables_text:
                text_content += "\n\n--- Tables ---\n" + "\n".join(tables_text)
            
            # 2. Extract Images (Multimodal)
            image_analysis_text = ""
            # Only proceed if we have Image library valid
            if Image:
                try:
                    images_found = []
                    # Scan for images in relationships
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            images_found.append(rel.target_part.blob)
                    
                    if images_found:
                        print(f"👁️ Found {len(images_found)} images in DOCX. Filtering and Analyzing via NVIDIA...")
                        
                        async def analyze_single_image_nvidia(img_bytes, idx):
                            try:
                                img = Image.open(io.BytesIO(img_bytes))
                                
                                # OPTIMIZATION 1: Skip Small Images (Icons/Logos)
                                width, height = img.size
                                if width < 100 or height < 100:
                                    # print(f"Skipping small image {idx+1} ({width}x{height})")
                                    return ""

                                await mistral_limiter.wait_for_slot()
                                
                                # Convert to base64
                                base64_img = _optimize_and_encode(img)
                                
                                headers = {
                                    "Authorization": f"Bearer {NVIDIA_KEY}",
                                    "Content-Type": "application/json",
                                    "Accept": "application/json"
                                }
                                
                                payload = {
                                    "model": "mistralai/mistral-large-3-675b-instruct-2512",
                                    "messages": [{
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": "Describe this image from a document in detail. Transcribe any text visible. Be concise."},
                                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_img}"}}
                                        ]
                                    }],
                                    "max_tokens": 1024,
                                    "temperature": 0.15
                                }
                                
                                async with httpx.AsyncClient() as http_client:
                                    resp = await http_client.post(NVIDIA_API_URL, headers=headers, json=payload, timeout=30.0)
                                    if resp.status_code == 200:
                                        data = resp.json()
                                        content = data['choices'][0]['message']['content']
                                        return f"\n\n## [Image {idx+1} Description]\n{content}\n"
                                    else:
                                        logger.error(f"NVIDIA API Error {resp.status_code}: {resp.text}")
                                        return f"\n[Image {idx+1} Analysis Failed: {resp.status_code}]"

                            except Exception as e:
                                logger.error(f"Image {idx+1} analysis failed: {e}")
                                return f"\n[Image {idx+1} Analysis Error]"

                        # Run parallel analysis
                        tasks = [analyze_single_image_nvidia(img, i) for i, img in enumerate(images_found)]
                        results = await asyncio.gather(*tasks)
                        # Filter out empty strings (skipped images)
                        valid_results = [r for r in results if r]
                        image_analysis_text = "".join(valid_results)
                        
                        if len(valid_results) < len(images_found):
                            print(f"⚡ Optimization: Processed {len(valid_results)}/{len(images_found)} images (skipped {len(images_found)-len(valid_results)} small ones).")

                except Exception as img_e:
                    logger.error(f"Image extraction failed: {img_e}")
                    image_analysis_text = f"\n[Image Extraction Error: {img_e}]"

            text = text_content + image_analysis_text

        else:
            # MD, TXT
            text = content.decode('utf-8', errors='ignore')
            
        return f"User Context: {user_prompt}\n\nDocument Content:\n{text}"
    except Exception as e:
        return f"Error processing text document: {str(e)}"

