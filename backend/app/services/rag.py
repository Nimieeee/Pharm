"""
RAG (Retrieval-Augmented Generation) Service
Handles document processing and vector search using LangChain + Supabase pgvector
"""

import os
import tempfile
from typing import List, Dict, Any, Optional, Tuple
from uuid import UUID
from supabase import Client
import re
from typing import List, Dict, Any, Optional, Tuple
from uuid import UUID
import numpy as np

# Additional imports for document processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.core.config import settings
from app.models.document import DocumentChunk, DocumentChunkCreate, DocumentUploadResponse


class SimpleDocument:
    """Simple document class to replace LangChain Document"""
    def __init__(self, page_content: str, metadata: Dict[str, Any] = None):
        self.page_content = page_content
        self.metadata = metadata or {}


class SimpleTextSplitter:
    """Simple text splitter to replace LangChain RecursiveCharacterTextSplitter"""
    
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 300):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = ["\n\n", "\n", ". ", " ", ""]
    
    def split_documents(self, documents: List[SimpleDocument]) -> List[SimpleDocument]:
        """Split documents into chunks"""
        chunks = []
        for doc in documents:
            text_chunks = self._split_text(doc.page_content)
            for chunk_text in text_chunks:
                chunk = SimpleDocument(
                    page_content=chunk_text,
                    metadata=doc.metadata.copy()
                )
                chunks.append(chunk)
        return chunks
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Try to find a good split point
            split_point = end
            for separator in self.separators:
                last_sep = text.rfind(separator, start, end)
                if last_sep > start:
                    split_point = last_sep + len(separator)
                    break
            
            chunks.append(text[start:split_point])
            start = split_point - self.chunk_overlap if self.chunk_overlap > 0 else split_point
        
        return [chunk.strip() for chunk in chunks if chunk.strip()]


class RAGService:
    """RAG service for document processing and retrieval using direct Supabase integration"""
    
    def __init__(self, db: Client):
        self.db = db
        self.text_splitter = SimpleTextSplitter(
            chunk_size=1500,
            chunk_overlap=300
        )
        self.mistral_api_key = os.getenv("MISTRAL_API_KEY")
        self.mistral_base_url = "https://api.mistral.ai/v1"
        print("✅ RAG service initialized with Mistral embeddings")
    
    def _generate_embedding(self, text: str) -> Optional[List[float]]:
        """Generate embedding using Mistral API"""
        try:
            if not text.strip():
                return None
            
            if not self.mistral_api_key:
                print("⚠️ Mistral API key not configured, using fallback")
                return self._generate_fallback_embedding(text)
            
            import httpx
            
            # Use Mistral's embedding model
            payload = {
                "model": "mistral-embed",
                "input": [text]
            }
            
            with httpx.Client() as client:
                response = client.post(
                    f"{self.mistral_base_url}/embeddings",
                    headers={
                        "Authorization": f"Bearer {self.mistral_api_key}",
                        "Content-Type": "application/json"
                    },
                    json=payload,
                    timeout=30.0
                )
                
                if response.status_code == 200:
                    result = response.json()
                    if result.get("data") and len(result["data"]) > 0:
                        embedding = result["data"][0]["embedding"]
                        print(f"✅ Generated Mistral embedding: {len(embedding)} dimensions")
                        return embedding
                    else:
                        print("❌ No embedding data in response")
                        return self._generate_fallback_embedding(text)
                else:
                    print(f"❌ Mistral API error: {response.status_code}")
                    return self._generate_fallback_embedding(text)
            
        except Exception as e:
            print(f"❌ Error generating Mistral embedding: {e}")
            return self._generate_fallback_embedding(text)
    
    def _generate_fallback_embedding(self, text: str) -> Optional[List[float]]:
        """Generate fallback hash-based embedding"""
        try:
            import hashlib
            hash_obj = hashlib.sha256(text.encode())
            hash_bytes = hash_obj.digest()
            
            # Convert to 1024-dimensional vector (Mistral embed dimension)
            embedding = []
            for i in range(1024):
                byte_val = hash_bytes[i % len(hash_bytes)]
                embedding.append((byte_val / 255.0) * 2.0 - 1.0)
            
            return embedding
        except Exception as e:
            print(f"❌ Error generating fallback embedding: {e}")
            return None
    
    async def process_uploaded_file(
        self, 
        file_content: bytes, 
        filename: str, 
        conversation_id: UUID, 
        user_id: UUID
    ) -> DocumentUploadResponse:
        """Process uploaded file and store chunks using LangChain"""
        try:
            # Save file temporarily
            file_extension = filename.lower().split('.')[-1]
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Load document
                documents = self._load_document(tmp_file_path, filename)
                
                if not documents:
                    return DocumentUploadResponse(
                        success=False,
                        message=f"Could not process file: {filename}",
                        chunk_count=0
                    )
                
                # Split into chunks
                chunks = self.text_splitter.split_documents(documents)
                
                if not chunks:
                    return DocumentUploadResponse(
                        success=False,
                        message=f"No content found in file: {filename}",
                        chunk_count=0
                    )
                
                # Add metadata to chunks
                for chunk in chunks:
                    chunk.metadata.update({
                        "user_id": str(user_id),
                        "conversation_id": str(conversation_id),
                        "filename": filename,
                        "file_type": file_extension
                    })
                
                # Store chunks using direct database operations
                success_count = 0
                for chunk in chunks:
                    if await self._process_and_store_chunk(
                        chunk, filename, conversation_id, user_id
                    ):
                        success_count += 1
                
                return DocumentUploadResponse(
                    success=success_count > 0,
                    message=f"Processed {success_count} chunks from {filename}",
                    chunk_count=success_count
                )
                
            finally:
                # Clean up temporary file
                os.unlink(tmp_file_path)
                
        except Exception as e:
            return DocumentUploadResponse(
                success=False,
                message=f"Error processing file: {str(e)}",
                chunk_count=0
            )
    
    async def _process_and_store_chunk(
        self, 
        chunk: SimpleDocument, 
        filename: str, 
        conversation_id: UUID, 
        user_id: UUID
    ) -> bool:
        """Process chunk and store in database"""
        try:
            # Generate embedding
            embedding = self._generate_embedding(chunk.page_content)
            if not embedding:
                return False
            
            # Determine embedding version
            embedding_version = "mistral-v1" if self.mistral_api_key else "hash-v1"
            
            # Prepare metadata
            metadata = {
                "filename": filename,
                "page": chunk.metadata.get("page", 0),
                "source": chunk.metadata.get("source", filename),
                "file_type": chunk.metadata.get("file_type", "unknown"),
                "user_id": str(user_id),
                "conversation_id": str(conversation_id)
            }
            
            # Store in database
            chunk_data = {
                "conversation_id": str(conversation_id),
                "user_id": str(user_id),
                "content": chunk.page_content,
                "embedding": embedding,
                "embedding_version": embedding_version,
                "metadata": metadata
            }
            
            result = self.db.table("document_chunks").insert(chunk_data).execute()
            return len(result.data) > 0
            
        except Exception as e:
            print(f"❌ Error storing chunk: {e}")
            return False
    
    def _load_document(self, file_path: str, filename: str) -> List[SimpleDocument]:
        """Load document based on file type"""
        try:
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self._process_pdf(file_path, filename)
                
            elif file_extension in ['txt', 'md']:
                return self._process_text(file_path, filename)
                
            elif file_extension == 'docx' and DOCX_AVAILABLE:
                return self._process_docx(file_path, filename)
                
            else:
                print(f"❌ Unsupported file type: {file_extension}")
                return []
                
        except Exception as e:
            print(f"❌ Error loading document '{filename}': {e}")
            return []
    
    def _process_pdf(self, file_path: str, filename: str) -> List[SimpleDocument]:
        """Process PDF files"""
        try:
            import PyPDF2
            documents = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        doc = SimpleDocument(
                            page_content=text,
                            metadata={
                                "source": filename,
                                "page": page_num + 1,
                                "file_type": "pdf"
                            }
                        )
                        documents.append(doc)
            
            return documents
        except Exception as e:
            print(f"❌ Error processing PDF '{filename}': {e}")
            return []
    
    def _process_text(self, file_path: str, filename: str) -> List[SimpleDocument]:
        """Process text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
            if content.strip():
                doc = SimpleDocument(
                    page_content=content,
                    metadata={
                        "source": filename,
                        "file_type": "text"
                    }
                )
                return [doc]
            
            return []
        except Exception as e:
            print(f"❌ Error processing text file '{filename}': {e}")
            return []
    
    def _process_docx(self, file_path: str, filename: str) -> List[SimpleDocument]:
        """Process DOCX files"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs and tables
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            if not text_content:
                return []
            
            full_text = '\n\n'.join(text_content)
            document = SimpleDocument(
                page_content=full_text,
                metadata={"source": filename, "file_type": "docx"}
            )
            
            return [document]
            
        except Exception as e:
            print(f"❌ Error processing DOCX '{filename}': {e}")
            return []
    

    

    
    async def search_similar_chunks(
        self, 
        query: str, 
        conversation_id: UUID, 
        user_id: UUID, 
        max_results: int = 10,
        similarity_threshold: float = 0.5
    ) -> List[DocumentChunk]:
        """Search for similar document chunks using direct database operations"""
        try:
            # Generate query embedding
            query_embedding = self._generate_embedding(query)
            if not query_embedding:
                print("⚠️ No query embedding generated, using recent chunks")
                return await self._get_recent_chunks(conversation_id, user_id, max_results)
            
            # Determine embedding dimension
            embedding_dim = len(query_embedding)
            print(f"🔍 Searching with {embedding_dim}-dimensional embedding")
            
            # Use appropriate function based on dimension
            if embedding_dim == 1024:
                function_name = 'match_documents_with_user_isolation_v2'
            else:
                function_name = 'match_documents_with_user_isolation'
            
            # Search using database function
            result = self.db.rpc(
                function_name,
                {
                    'query_embedding': query_embedding,
                    'conversation_uuid': str(conversation_id),
                    'user_session_uuid': str(user_id),
                    'match_threshold': similarity_threshold,
                    'match_count': max_results
                }
            ).execute()
            
            chunks = []
            for row in result.data or []:
                chunk = DocumentChunk(
                    id=row['id'],
                    conversation_id=conversation_id,
                    content=row['content'],
                    metadata=row['metadata'],
                    similarity=row['similarity'],
                    created_at=row.get('created_at', '2024-01-01T00:00:00Z')
                )
                chunks.append(chunk)
            
            print(f"✅ Found {len(chunks)} similar chunks")
            return chunks
            
        except Exception as e:
            print(f"❌ Error searching chunks: {e}")
            # Fallback to recent chunks
            return await self._get_recent_chunks(conversation_id, user_id, max_results)
    

    
    async def _get_recent_chunks(
        self, 
        conversation_id: UUID, 
        user_id: UUID, 
        limit: int = 10
    ) -> List[DocumentChunk]:
        """Get recent chunks as fallback"""
        try:
            result = self.db.table("document_chunks").select(
                "id, content, metadata, created_at"
            ).eq("conversation_id", str(conversation_id)).eq(
                "user_id", str(user_id)
            ).order("created_at", desc=True).limit(limit).execute()
            
            chunks = []
            for row in result.data or []:
                chunk = DocumentChunk(
                    id=row['id'],
                    conversation_id=conversation_id,
                    content=row['content'],
                    metadata=row['metadata'],
                    similarity=0.5,  # Default similarity
                    created_at=row['created_at']
                )
                chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            print(f"❌ Error getting recent chunks: {e}")
            return []
    
    async def get_all_conversation_chunks(
        self, 
        conversation_id: UUID, 
        user_id: UUID
    ) -> List[DocumentChunk]:
        """Get all chunks for a conversation"""
        try:
            result = self.db.table("document_chunks").select(
                "id, content, metadata, created_at"
            ).eq("conversation_id", str(conversation_id)).eq(
                "user_id", str(user_id)
            ).order("created_at").execute()
            
            chunks = []
            for row in result.data or []:
                chunk = DocumentChunk(
                    id=row['id'],
                    conversation_id=conversation_id,
                    content=row['content'],
                    metadata=row['metadata'],
                    similarity=0.5,
                    created_at=row['created_at']
                )
                chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            print(f"❌ Error getting all chunks: {e}")
            return []
    
    async def get_conversation_context(
        self, 
        query: str, 
        conversation_id: UUID, 
        user_id: UUID,
        max_chunks: int = 20
    ) -> str:
        """Get relevant context for a query"""
        try:
            # Search for relevant chunks
            chunks = await self.search_similar_chunks(
                query, conversation_id, user_id, max_chunks, 0.1
            )
            
            if not chunks:
                return ""
            
            # Organize by document
            documents = {}
            for chunk in chunks:
                filename = chunk.metadata.get("filename", "Unknown")
                if filename not in documents:
                    documents[filename] = []
                documents[filename].append(chunk.content)
            
            # Build context
            context_parts = []
            
            if len(documents) > 1:
                context_parts.append("=== DOCUMENT OVERVIEW ===")
                for filename in documents:
                    context_parts.append(f"📄 {filename} - {len(documents[filename])} sections")
                context_parts.append("")
            
            # Add content
            for filename, contents in documents.items():
                if len(documents) > 1:
                    context_parts.append(f"=== {filename} ===")
                
                for content in contents:
                    if content.strip():
                        context_parts.append(content.strip())
                
                if len(documents) > 1:
                    context_parts.append("")
            
            return "\n\n".join(context_parts)
            
        except Exception as e:
            print(f"❌ Error getting context: {e}")
            return ""
    
    async def delete_conversation_documents(
        self, 
        conversation_id: UUID, 
        user_id: UUID
    ) -> bool:
        """Delete all documents for a conversation"""
        try:
            result = self.db.table("document_chunks").delete().eq(
                "conversation_id", str(conversation_id)
            ).eq("user_id", str(user_id)).execute()
            
            return True
            
        except Exception as e:
            print(f"❌ Error deleting documents: {e}")
            return False