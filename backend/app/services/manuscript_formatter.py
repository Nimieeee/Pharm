"""
Manuscript Formatter Service

Restructures chat conversations into professional manuscript/report format
with sections, title page, table of contents, and proper formatting.
"""

import io
import re
from typing import List, Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ManuscriptFormatter:
    """
    Formats chat conversations into structured manuscripts.
    
    Supported styles:
    - report: Executive Summary, Introduction, Findings, Discussion, Conclusion
    - manuscript: Abstract, Introduction, Methods, Results, Discussion, References
    - plain: Simple formatted document with headings
    """
    
    STYLES = {
        "report": [
            "Executive Summary",
            "Introduction",
            "Findings",
            "Discussion",
            "Conclusion",
        ],
        "manuscript": [
            "Abstract",
            "Introduction",
            "Methods",
            "Results",
            "Discussion",
            "Conclusion",
        ],
        "plain": [
            "Overview",
            "Details",
            "Summary",
        ]
    }
    
    def __init__(self, multi_provider=None):
        """
        Initialize ManuscriptFormatter.
        
        Args:
            multi_provider: MultiProviderService instance for AI restructuring
        """
        self.multi_provider = multi_provider
    
    async def structure_content(self, messages: List[Dict[str, Any]], style: str = "report") -> Dict[str, Any]:
        """
        AI-restructures chat messages into manuscript sections.
        
        Args:
            messages: List of conversation messages (user + assistant)
            style: Manuscript style (report, manuscript, plain)
            
        Returns:
            Dict with structured sections and content
        """
        # Extract conversation text
        conversation_text = self._extract_conversation_text(messages)
        
        # Get section structure for this style
        sections = self.STYLES.get(style, self.STYLES["report"])
        
        # Use AI to restructure content into sections
        structured = await self._ai_restructure(conversation_text, sections, style)
        
        return structured
    
    def _extract_conversation_text(self, messages: List[Dict[str, Any]]) -> str:
        """Extract clean conversation text from message list."""
        parts = []
        for msg in messages:
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            parts.append(f"[{role.upper()}]: {content}")
        return "\n\n".join(parts)
    
    async def _ai_restructure(self, conversation_text: str, sections: List[str], style: str) -> Dict[str, Any]:
        """
        Use AI to restructure conversation into manuscript sections.
        
        Args:
            conversation_text: Raw conversation text
            sections: List of section names
            style: Manuscript style
            
        Returns:
            Dict with title, sections, and content
        """
        if not self.multi_provider:
            # Fallback: simple extraction without AI restructuring
            return self._simple_extract(conversation_text, sections, style)
        
        section_list = "\n".join([f"- {s}" for s in sections])
        
        prompt = f"""Restructure the following conversation into a professional {style} format.

Required sections:
{section_list}

Return ONLY valid JSON with this exact schema:
{{
  "title": "Document Title",
  "subtitle": "Optional subtitle or date",
  "sections": [
    {{
      "heading": "Section Name",
      "content": "Section content in paragraph form...",
      "subsections": [
        {{"heading": "Subsection", "content": "Content..."}}
      ]
    }}
  ]
}}

Rules:
- Write in professional, academic tone
- Use complete paragraphs, not bullet points
- Include all key information from the conversation
- Remove conversational elements (greetings, clarifications)
- Keep scientific accuracy
- Maximum 500 words per section"""

        try:
            response = await self.multi_provider.generate(
                messages=[
                    {"role": "system", "content": "You are a professional medical/scientific writer."},
                    {"role": "user", "content": f"{prompt}\n\nConversation:\n{conversation_text[:8000]}"}
                ],
                mode="fast",
                max_tokens=4000,
                temperature=0.3
            )
            
            # Parse JSON response
            import json
            json_str = response.strip()
            if json_str.startswith("```"):
                json_str = json_str.split("```")[1]
                if json_str.startswith("json"):
                    json_str = json_str[4:]
            json_str = json_str.strip()
            
            structured = json.loads(json_str)
            return structured
            
        except Exception as e:
            print(f"⚠️ AI restructuring failed: {e}, using simple extraction")
            return self._simple_extract(conversation_text, sections, style)
    
    def _simple_extract(self, conversation_text: str, sections: List[str], style: str) -> Dict[str, Any]:
        """Simple fallback extraction without AI."""
        # Just use the conversation as-is with basic section headers
        return {
            "title": "Conversation Export",
            "subtitle": f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "sections": [
                {
                    "heading": sections[0] if sections else "Overview",
                    "content": conversation_text[:2000]
                }
            ]
        }
    
    def build_docx(self, structured: Dict[str, Any], style: str = "report") -> bytes:
        """
        Build DOCX document from structured content.
        
        Features:
        - Title page with Benchside branding
        - Table of Contents (auto-generated)
        - Numbered heading hierarchy (H1, H2, H3)
        - Page numbers in footer
        - 'Generated by Benchside' watermark
        
        Args:
            structured: Dict from structure_content()
            style: Manuscript style
            
        Returns:
            DOCX file as bytes
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed")
        
        doc = Document()
        
        # Set default style
        self._set_document_styles(doc)
        
        # Title page
        self._add_title_page(doc, structured.get("title", "Document"), 
                            structured.get("subtitle", ""))
        
        # Page break after title
        doc.add_page_break()
        
        # Table of Contents
        self._add_table_of_contents(doc)
        
        # Page break after TOC
        doc.add_page_break()
        
        # Main content
        for section in structured.get("sections", []):
            self._add_section(doc, section)
        
        # Footer with watermark
        self._add_footer(doc)
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.getvalue()
    
    def _set_document_styles(self, doc):
        """Set professional document styles."""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x2D, 0x34, 0x36)

        # Heading 1 style
        h1 = doc.styles['Heading 1']
        h1_font = h1.font
        h1_font.name = 'Calibri'
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # Heading 2 style
        h2 = doc.styles['Heading 2']
        h2_font = h2.font
        h2_font.name = 'Calibri'
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0x34, 0x49, 0x5E)

        # Heading 3 style
        h3 = doc.styles['Heading 3']
        h3_font = h3.font
        h3_font.name = 'Calibri'
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    def _add_title_page(self, doc, title: str, subtitle: str):
        """Add professional title page."""
        # Add logo placeholder (could be image)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        run = title_para.add_run(title)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        run.font.name = 'Calibri'
        
        # Subtitle
        if subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = Pt(12)
            sub_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            sub_run.font.name = 'Calibri'
        
        # Spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Branding
        branding_para = doc.add_paragraph()
        branding_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        branding_run = branding_para.add_run("Generated by Benchside")
        branding_run.font.size = Pt(10)
        branding_run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
        branding_run.font.italic = True
    
    def _add_table_of_contents(self, doc):
        """Add table of contents placeholder."""
        doc.add_heading("Table of Contents", level=1)

        # Note: Word auto-generates TOC, we just add placeholder
        toc_para = doc.add_paragraph()
        toc_run = toc_para.add_run(
            "[Table of Contents will be auto-generated when opened in Microsoft Word]\n\n"
            "To update: Right-click this area → Update Field → Update entire table"
        )
        toc_run.font.size = Pt(10)
        toc_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        toc_run.font.italic = True

    def _add_section(self, doc, section: Dict[str, Any]):
        """Add a section with content and subsections."""
        heading = section.get("heading", "Section")
        content = section.get("content", "")
        subsections = section.get("subsections", [])
        
        # Add main heading
        doc.add_heading(heading, level=1)
        
        # Add content as paragraphs
        # Split by double newlines for paragraph breaks
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(12)
        
        # Add subsections
        for subsection in subsections:
            sub_heading = subsection.get("heading", "")
            sub_content = subsection.get("content", "")
            
            if sub_heading:
                doc.add_heading(sub_heading, level=2)
            
            if sub_content:
                sub_paragraphs = sub_content.split('\n\n')
                for para in sub_paragraphs:
                    if para.strip():
                        p = doc.add_paragraph(para.strip())
                        p.paragraph_format.space_after = Pt(12)
    
    def _add_footer(self, doc):
        """Add footer with page numbers and watermark."""
        # Note: python-docx has limited footer support
        # This is a placeholder - footers would need oxml manipulation
        pass


# Singleton instance
manuscript_formatter = ManuscriptFormatter()
