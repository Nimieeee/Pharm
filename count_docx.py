import sys
import docx
import tiktoken

def count_words_and_tokens(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also get tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    text = '\n'.join(full_text)
    words = len(text.split())
    
    # Use standard cl100k_base encoding for GPT-4/Llama token approximation
    enc = tiktoken.get_encoding("cl100k_base")
    tokens = len(enc.encode(text))
    
    print(f"File: {filename}")
    print(f"Words: {words}")
    print(f"Approximate Tokens: {tokens}")
    print(f"Character length: {len(text)}")

if __name__ == '__main__':
    count_words_and_tokens(sys.argv[1])
